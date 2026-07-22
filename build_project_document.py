from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.section import WD_SECTION
from pathlib import Path

OUT = Path("Documento_inicial_SaaS_restaurantes.docx")
NAVY = RGBColor(31, 77, 120)
BLUE = RGBColor(46, 116, 181)
GRAY = RGBColor(90, 98, 108)
LIGHT = "E8EEF5"

def font(run, size=11, bold=False, color=None, italic=False):
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color

def shade(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = tcPr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tcPr.append(shd)
    shd.set(qn("w:fill"), fill)

def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = tcPr.first_child_found_in("w:tcMar")
    if tcMar is None:
        tcMar = OxmlElement("w:tcMar")
        tcPr.append(tcMar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tcMar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tcMar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")

def set_table_widths(table, widths):
    table.autofit = False
    tblPr = table._tbl.tblPr
    tblW = tblPr.find(qn("w:tblW"))
    if tblW is None:
        tblW = OxmlElement("w:tblW")
        tblPr.append(tblW)
    tblW.set(qn("w:w"), str(sum(widths)))
    tblW.set(qn("w:type"), "dxa")
    tblInd = tblPr.find(qn("w:tblInd"))
    if tblInd is None:
        tblInd = OxmlElement("w:tblInd")
        tblPr.append(tblInd)
    tblInd.set(qn("w:w"), "120")
    tblInd.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid): grid.remove(child)
    for w in widths:
        col = OxmlElement("w:gridCol"); col.set(qn("w:w"), str(w)); grid.append(col)
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            tcW = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
            if tcW is None:
                tcW = OxmlElement("w:tcW"); cell._tc.get_or_add_tcPr().append(tcW)
            tcW.set(qn("w:w"), str(widths[i])); tcW.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(.5)
    p.paragraph_format.first_line_indent = Inches(-.25)
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.line_spacing = 1.15
    font(p.add_run(text))
    return p

def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.left_indent = Inches(.5)
    p.paragraph_format.first_line_indent = Inches(-.25)
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.line_spacing = 1.15
    font(p.add_run(text))
    return p

doc = Document()
sec = doc.sections[0]
sec.page_width, sec.page_height = Inches(8.5), Inches(11)
sec.top_margin = sec.bottom_margin = sec.left_margin = sec.right_margin = Inches(1)
sec.header_distance = sec.footer_distance = Inches(.492)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Calibri"; normal.font.size = Pt(11); normal.font.color.rgb = RGBColor(0,0,0)
normal.paragraph_format.space_after = Pt(6); normal.paragraph_format.line_spacing = 1.10
for name, size, before, after, color in [
    ("Heading 1",16,16,8,BLUE),("Heading 2",13,12,6,BLUE),("Heading 3",12,8,4,NAVY)]:
    s=styles[name]; s.font.name="Calibri"; s.font.size=Pt(size); s.font.bold=True; s.font.color.rgb=color
    s.paragraph_format.space_before=Pt(before); s.paragraph_format.space_after=Pt(after); s.paragraph_format.keep_with_next=True

# Header and footer
h = sec.header.paragraphs[0]
h.alignment = WD_ALIGN_PARAGRAPH.RIGHT
font(h.add_run("Especificación inicial | SaaS para restaurantes"), 9, color=GRAY)
f = sec.footer.paragraphs[0]
f.alignment = WD_ALIGN_PARAGRAPH.CENTER
font(f.add_run("Documento de trabajo • versión 0.1 • julio de 2026"), 9, color=GRAY)

# Cover
p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(72); p.paragraph_format.space_after=Pt(10)
font(p.add_run("DOCUMENTO INICIAL DEL PROYECTO"), 11, True, BLUE)
p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(8)
font(p.add_run("Plataforma SaaS para restaurantes"), 28, True, NAVY)
p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(28)
font(p.add_run("Pedidos móviles, operación, control financiero e informes en una plataforma multiempresa."), 14, color=GRAY)

t=doc.add_table(rows=4, cols=2); t.alignment=WD_TABLE_ALIGNMENT.LEFT; t.style="Table Grid"
rows=[("Tipo de producto","Software como servicio (SaaS) multiempresa"),("Modelo inicial","1 restaurante, 1 sede y 2 conexiones móviles simultáneas"),("Monetización","Suscripción base más conexiones, sedes y módulos adicionales"),("Estado","Definición inicial para validación y desarrollo del MVP")]
for i,(a,b) in enumerate(rows):
    t.cell(i,0).text=a; t.cell(i,1).text=b; shade(t.cell(i,0),LIGHT)
    for r in t.cell(i,0).paragraphs[0].runs: font(r,10.5,True,NAVY)
    for r in t.cell(i,1).paragraphs[0].runs: font(r,10.5)
set_table_widths(t,[2700,6660])

doc.add_page_break()
doc.add_heading("1. Decisión fundacional", level=1)
p=doc.add_paragraph()
font(p.add_run("La solución se diseñará desde su primera versión como un SaaS multiempresa."),11,True,NAVY)
font(p.add_run(" Cada cliente tendrá datos aislados, una suscripción, límites de uso y módulos contratados. La arquitectura debe permitir crecer a varias sedes y marcas sin reconstruir el sistema."))

doc.add_heading("2. Objetivo del producto", level=1)
p=doc.add_paragraph("Centralizar la operación diaria y el control administrativo de restaurantes mediante aplicaciones móviles y un panel web, conectando pedidos, cocina, caja, inventario, gastos e informes.")

doc.add_heading("3. Estructura de la plataforma", level=1)
for x in ["Plataforma: administra clientes, planes, módulos, pagos y soporte.","Cliente o empresa: titular de la suscripción y propietario de sus datos.","Restaurante o marca: unidad comercial registrada por el cliente.","Sede: ubicación operativa con mesas, caja, inventario y personal.","Usuario: persona con identidad, rol y permisos.","Dispositivo y sesión: acceso móvil sujeto al límite simultáneo contratado."]:
    add_bullet(doc,x)

doc.add_heading("4. Oferta inicial", level=1)
doc.add_heading("4.1 Plan base", level=2)
for x in ["Un restaurante y una sede.","Dos conexiones móviles simultáneas.","Panel administrativo web.","Usuarios, roles y permisos básicos.","Menú, categorías, productos, variantes y adicionales.","Mesas, pedidos en salón y pedidos para llevar.","Pantalla de cocina y estados de preparación.","Caja, medios de pago, apertura y cierre.","Registro básico de gastos e ingresos no operativos.","Informes básicos de ventas, caja y operación."]:
    add_bullet(doc,x)

doc.add_heading("4.2 Complementos cobrados por separado", level=2)
for x in ["Conexiones móviles simultáneas adicionales.","Sedes o restaurantes adicionales.","Inventario avanzado, recetas, costos y desperdicios.","Compras y proveedores.","Facturación electrónica.","Pedidos por código QR, reservas y domicilios.","Fidelización, promociones e historial de clientes.","Integraciones contables, de pago o de entrega.","Informes financieros avanzados y soporte prioritario."]:
    add_bullet(doc,x)

doc.add_heading("5. Regla de conexiones móviles", level=1)
p=doc.add_paragraph("El límite se aplicará a conexiones móviles simultáneas, no a la cantidad de empleados registrados. Una cuenta podrá crear varios usuarios, pero solo el número contratado podrá mantener sesiones móviles activas al mismo tiempo.")
for x in ["Mostrar las sesiones y dispositivos activos.","Permitir a un usuario autorizado cerrar una sesión anterior.","Bloquear la conexión que exceda el límite sin afectar pedidos ya registrados.","Ofrecer una ampliación del plan y registrar el intento en auditoría.","Definir una política segura para sesiones abandonadas, cambio de dispositivo y pérdida del teléfono."]:
    add_bullet(doc,x)

doc.add_heading("6. Módulos funcionales", level=1)
modules=[
 ("Acceso y seguridad","Registro, autenticación, recuperación, roles, permisos y auditoría."),
 ("Configuración","Datos fiscales, sedes, impuestos, monedas, horarios y parámetros."),
 ("Menú","Categorías, productos, precios, modificadores, disponibilidad y promociones."),
 ("Pedidos","Mesas, salón, para llevar, observaciones, división y traslado de cuentas."),
 ("Cocina","Comandas por estación, prioridades, tiempos y estados de preparación."),
 ("Caja","Turnos, cobros, pagos combinados, propinas, devoluciones y arqueos."),
 ("Finanzas operativas","Ingresos, gastos, soportes, categorías y responsables."),
 ("Inventario","Ingredientes, recetas, compras, existencias, conteos y desperdicios."),
 ("Informes","Ventas, caja, costos, rentabilidad, inventario y productividad."),
 ("Suscripciones","Planes, módulos, límites, pagos, pruebas, suspensión y reactivación."),
]
t=doc.add_table(rows=1,cols=2); t.style="Table Grid"; t.alignment=WD_TABLE_ALIGNMENT.LEFT
t.cell(0,0).text="Módulo"; t.cell(0,1).text="Responsabilidad principal"
for c in t.rows[0].cells: shade(c,LIGHT)
for a,b in modules:
    cells=t.add_row().cells; cells[0].text=a; cells[1].text=b
for row in t.rows:
    for j,c in enumerate(row.cells):
        for r in c.paragraphs[0].runs: font(r,10, row is t.rows[0] or j==0, NAVY if (row is t.rows[0] or j==0) else None)
set_table_widths(t,[2400,6960])

doc.add_heading("7. Informes mínimos del MVP", level=1)
for x in ["Ventas por fecha, sede, producto, categoría, mesero y canal.","Cantidad de pedidos, clientes atendidos y valor promedio por pedido.","Productos más y menos vendidos; horas y días de mayor demanda.","Resumen de caja, medios de pago, diferencias, anulaciones y descuentos.","Ingresos, gastos y resultado operativo estimado.","Inventario actual, faltantes y costo teórico cuando el módulo esté contratado."]:
    add_bullet(doc,x)
p=doc.add_paragraph(); font(p.add_run("Criterio contable: "),11,True,NAVY); font(p.add_run("el sistema diferenciará ventas, flujo de caja y utilidad para evitar informes engañosos."))

doc.add_heading("8. Administración del SaaS", level=1)
for x in ["Panel interno independiente del panel del restaurante.","Gestión de clientes, pruebas, planes, módulos, límites y promociones.","Control de pagos, periodos de gracia, suspensión y reactivación.","Métricas de uso, actividad, errores y consumo por cuenta.","Soporte con acceso controlado, temporal y auditado, sin conocer contraseñas.","Conservación de datos ante falta de pago; nunca borrar automáticamente información operativa."]:
    add_bullet(doc,x)

doc.add_heading("9. Estados de la suscripción", level=1)
for x in ["Prueba", "Activa", "Pago pendiente", "Periodo de gracia", "Suspendida", "Cancelada", "Archivada"]: add_bullet(doc,x)

doc.add_heading("10. Requisitos no funcionales", level=1)
for x in ["Aislamiento estricto de datos entre clientes.","Cifrado, copias de seguridad y recuperación ante fallos.","Auditoría de cobros, anulaciones, descuentos, cierres y cambios administrativos.","Buen desempeño en teléfonos económicos y conexiones móviles inestables.","Escalabilidad para múltiples clientes, sedes y volúmenes de pedidos.","Configuración regional de moneda, zona horaria, impuestos e idioma.","Protección de datos personales y cumplimiento legal del país de operación."]:
    add_bullet(doc,x)

doc.add_heading("11. Alcance recomendado del MVP", level=1)
for x in ["Cuentas SaaS, suscripciones y aislamiento multiempresa.","Restaurante, sede, usuarios, roles y dos conexiones móviles.","Menú, mesas, pedidos y aplicación móvil del mesero.","Pantalla de cocina.","Caja, pagos, gastos y cierres.","Panel administrativo e informes básicos.","Panel interno para gestionar clientes, planes y módulos."]:
    add_number(doc,x)

doc.add_heading("12. Decisiones pendientes", level=1)
pending=[("País y marco fiscal","Define impuestos, moneda y facturación electrónica."),("Cliente objetivo","Restaurante independiente, cadena pequeña, bar, cafetería o comida rápida."),("Canales del MVP","Salón, para llevar, domicilio y/o QR."),("Operación sin Internet","Nivel de continuidad requerido cuando falle la conexión."),("Precios","Valor del plan base, conexión adicional, sede y módulos."),("Inventario","Básico dentro del plan o completamente adicional."),("Equipos","Android, iOS, tabletas, impresoras térmicas, cajón y lector."),("Periodo de prueba","Duración, límites y necesidad de medio de pago.")]
t=doc.add_table(rows=1,cols=2); t.style="Table Grid"; t.alignment=WD_TABLE_ALIGNMENT.LEFT
t.cell(0,0).text="Decisión"; t.cell(0,1).text="Impacto"
for c in t.rows[0].cells: shade(c,LIGHT)
for a,b in pending:
    cells=t.add_row().cells; cells[0].text=a; cells[1].text=b
for i,row in enumerate(t.rows):
    for j,c in enumerate(row.cells):
        for r in c.paragraphs[0].runs: font(r,10,i==0 or j==0,NAVY if i==0 or j==0 else None)
set_table_widths(t,[2700,6660])

doc.add_heading("13. Próxima etapa", level=1)
p=doc.add_paragraph("Validar las decisiones pendientes y convertir este documento en requisitos funcionales detallados. Después se definirán los flujos de usuario, arquitectura técnica, modelo de datos, prototipos de pantallas y plan de desarrollo por entregas.")

doc.core_properties.title="Plataforma SaaS para restaurantes - Documento inicial"
doc.core_properties.subject="Definición inicial del producto, modelo SaaS y alcance del MVP"
doc.core_properties.author="Equipo del proyecto"
doc.save(OUT)
print(OUT.resolve())
